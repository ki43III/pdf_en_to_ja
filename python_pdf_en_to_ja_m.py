import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
import nltk
import re # 正規表現モジュールをインポート
import io # 画像処理用
import threading
import queue
from docx import Document  # python-docx をインポート
from docx.shared import Inches  # 画像サイズ調整用
import glob
import os

# スレッド数
NUM_THREADS = 4

def translate_sentence(sentence_queue, result_queue):
    """翻訳スレッドで実行される関数"""
    translator = GoogleTranslator(source='en', target='ja')
    while True:
        sentence = sentence_queue.get()
        if sentence is None:
            break
        try:
            translated = translator.translate(sentence)
            result_queue.put((sentence, translated))
        except Exception as e:
            result_queue.put((sentence, str(e)))

# NLTKの文分割用データをダウンロード (初回のみ)
try:
    # punkt と punkt_tab が存在するか確認
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('tokenizers/punkt_tab') # punkt_tab も確認
except LookupError: # LookupError をキャッチするように変更
    print("NLTKの'punkt'および'punkt_tab'データをダウンロードします...")
    nltk.download('punkt')
    nltk.download('punkt_tab') # punkt_tab もダウンロード
    print("ダウンロード完了。")

print("ok")

# inフォルダ内の全PDFファイルを取得
input_dir = "./in"
pdf_files = glob.glob(os.path.join(input_dir, "*.pdf"))

if not pdf_files:
    print("inフォルダにPDFファイルが見つかりません。")
    exit()

for pdf_filename in pdf_files:
    print(f"=== {pdf_filename} の翻訳処理を開始 ===")
    output_filename = f"./out/output_{os.path.splitext(os.path.basename(pdf_filename))[0]}.docx"
    doc = fitz.open(pdf_filename)
    document = Document()

    # 日本語フォントの準備 (システムフォントパスを指定)
    font_path = "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc"
    font_name_ref = "hiragino" # PDF内でフォントを参照する名前

    # フォントの存在確認はここで行う (オプション)
    try:
        with open(font_path, 'rb') as f_check:
            pass
        print(f"フォントファイルが見つかりました: {font_path}")
    except FileNotFoundError:
        print(f"エラー: 指定されたフォントファイルが見つかりません: {font_path}")
        print("システムにインストールされている別の日本語フォントのパスを指定してください。")
        exit()
    except Exception as e:
         print(f"エラー: フォントファイルの読み込み中に問題が発生しました: {font_path}")
         print(f"詳細: {e}")
         exit()

    # ★★★ デバッグ: 実行時の PyMuPDF バージョンを表示 ★★★
    print(f"実行中の PyMuPDF バージョン: {fitz.__version__}")
    print(f"fitz モジュールの場所: {fitz.__file__}")

    # テキスト挿入用の設定
    start_point = fitz.Point(50, 70) # 左マージン
    right_margin = 50
    line_height = 14  # 基本の行の高さ
    max_line_spacing = 1.5  # 最大行間係数

    total_sentence_index = 1 # 全体での文番号を追跡

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        page_width = page.rect.width
        page_height = page.rect.height

        print(f"\n--- Processing Page {page_num + 1} ---")

        # --- 要素の抽出 (テキストと画像をdictから取得) ---
        all_elements = []
        image_data_list = [] # ページ内の画像データを保持 (get_images)
        try:
            # 1. get_text("dict") でテキストと画像のブロック情報を取得
            blocks_data = page.get_text("dict", flags=fitz.TEXTFLAGS_DICT & ~fitz.TEXT_INHIBIT_SPACES)["blocks"]
            for block in blocks_data:
                # bbox が存在し、有効か確認
                if "bbox" in block and len(block["bbox"]) == 4:
                    block["y0"] = block["bbox"][1] # ソート用にy座標を追加
                    all_elements.append(block)
                else:
                    print(f"  Warning: Block found without valid bbox on page {page_num + 1}. Type: {block.get('type', 'N/A')}. Skipping.")

            # y座標で要素をソート
            all_elements.sort(key=lambda x: x.get("y0", float('inf')))

            # 2. page.get_images() で画像データを別途取得 (後で使うため)
            img_list = page.get_images(full=True)
            image_data_list = [(img[0], doc.extract_image(img[0])) for img in img_list if doc.extract_image(img[0])] # (xref, image_dict) のリスト
            print(f"  Found {len([el for el in all_elements if el.get('type') == 0])} text blocks and {len([el for el in all_elements if el.get('type') == 1])} image blocks in dict.")
            print(f"  Found {len(image_data_list)} actual image data entries via get_images.")

        except Exception as extract_e:
            print(f"Error extracting elements from page {page_num + 1}: {extract_e}")
            # フォールバックはテキストのみを対象とする (従来通り)
            print("  Falling back to simple text extraction for this page.")
            text = page.get_text("text")
            text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
            text = re.sub(r' +', ' ', text).strip()
            sentences = nltk.sent_tokenize(text)
            all_elements = [{"type": "fallback_text", "sentences": sentences}]
            image_data_list = []

        # --- ソートされた要素を処理 ---
        image_data_iter = iter(image_data_list) # 画像データリストのイテレータ

        for element in all_elements:
            if element.get("type") == 0: # 通常のテキストブロック
                # --- テキストブロックの処理 ---
                block_text = ""
                try:
                    if "lines" not in element: continue # lines がなければスキップ
                    for line in element["lines"]:
                        if "spans" not in line: continue
                        for span in line["spans"]:
                            if "text" in span:
                                block_text += span["text"] + " "
                    block_text = block_text.strip()
                    block_text = re.sub(r' +', ' ', block_text)
                    if not block_text: continue

                    print(f"  Processing text block at original y0: {element.get('y0', 'N/A'):.2f} (len: {len(block_text)})")
                    sentences = nltk.sent_tokenize(block_text)

                    # スレッド間で共有するキュー
                    sentence_queue = queue.Queue()
                    result_queue = queue.Queue()

                    # 文をキューに追加
                    for sentence in sentences:
                        sentence_queue.put(sentence)

                    # 終了シグナルをキューに追加
                    for _ in range(NUM_THREADS):
                        sentence_queue.put(None)

                    # スレッドを起動
                    threads = []
                    for _ in range(NUM_THREADS):
                        thread = threading.Thread(target=translate_sentence, args=(sentence_queue, result_queue))
                        threads.append(thread)
                        thread.start()

                    # スレッドの終了を待機
                    for thread in threads:
                        thread.join()

                    # 結果を処理
                    while not result_queue.empty():
                        sentence, translated = result_queue.get()
                        if not sentence.strip(): continue
                        try:
                            print(f"    Translating sentence {total_sentence_index}...")
                            # Wordドキュメントに翻訳されたテキストを追加
                            document.add_paragraph(translated)
                            total_sentence_index += 1
                        except Exception as e:
                            print(f"    Error translating/inserting sentence {total_sentence_index}: {e}")
                            # エラーメッセージ挿入処理 (省略) ...
                            error_text = f"Sentence {total_sentence_index} (Error): {e}\nOriginal: {sentence}"
                            document.add_paragraph(error_text)
                            total_sentence_index += 1
                except Exception as block_proc_e:
                     print(f"  Error processing text block content: {block_proc_e}")

            elif element.get("type") == 1: # 画像ブロック (dict からの情報)
                # --- 画像ブロックの処理 (不安定な可能性あり) ---
                print(f"  Processing image block at original y0: {element.get('y0', 'N/A'):.2f}")
                try:
                    # image_data_list から次の画像データを取得 (順序が一致する前提)
                    xref, base_image = next(image_data_iter)
                    if not base_image:
                        print(f"    Warning: No image data found for assumed next image (xref: {xref}). Skipping.")
                        continue

                    image_bytes = base_image["image"]
                    img_w = base_image.get("width", 100)
                    img_h = base_image.get("height", 100)

                    # 画像を Word ドキュメントに追加
                    image_stream = io.BytesIO(image_bytes)
                    document.add_picture(image_stream, width=Inches(6))  # 幅を6インチに設定

                except StopIteration:
                    print("    Warning: No more image data available in image_data_list, but image block found in dict.")
                except Exception as e:
                    print(f"  Error processing image block: {e}")

            elif element.get("type") == "fallback_text":
                 # --- フォールバック処理 (ブロック抽出失敗時) ---
                 print("  Processing text using fallback method...")
                 sentences = element.get("sentences", [])
                 # (翻訳と挿入処理 - 上記テキストブロック内の処理と同様)
                 # ... (省略 - 上記の try...except ブロックを参考に実装) ...
                 for sentence in sentences:
                     if not sentence.strip(): continue
                     try:
                         print(f"    Translating sentence {total_sentence_index} (fallback)...")
                         translated = GoogleTranslator(source='en', target='ja').translate(sentence)
                         document.add_paragraph(translated)
                         total_sentence_index += 1
                     except Exception as e:
                         print(f"    Error translating/inserting sentence {total_sentence_index} (fallback): {e}")
                         error_text = f"Sentence {total_sentence_index} (Error): {e}\nOriginal: {sentence}"
                         document.add_paragraph(error_text)
                         total_sentence_index += 1

    # 作成した Word ドキュメントを保存
    try:
        document.save(output_filename)
        print(f"\n翻訳結果を {output_filename} に出力しました。")
    except Exception as e:
        print(f"\nエラー: Wordファイルの保存中にエラーが発生しました: {output_filename}")
        print(f"詳細: {e}")
    finally:
        # ドキュメントが開いているか確認してから閉じる
        if 'doc' in locals() and doc is not None and not doc.is_closed:
            doc.close() # 元のドキュメントも閉じる
